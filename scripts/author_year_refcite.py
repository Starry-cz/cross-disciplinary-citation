#!/usr/bin/env python3
"""将 DOCX 中精确匹配的著者—年份引文写为可点击的 Word REF 域。"""

import argparse
import os
import re
import shutil
import sys
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REF_TITLES = {"参考文献", "〖参考文献〗", "主要参考文献", "References", "Bibliography"}


def read_mapping(path):
    """读取精确引文到文后编号的映射，拒绝歧义和隐式修复。"""
    mapping = {}
    with open(path, "r", encoding="utf-8-sig") as handle:
        for line_no, raw in enumerate(handle, 1):
            line = raw.strip()
            if not line or line.startswith("#"):
                continue
            if line.count("=") != 1:
                raise ValueError(f"第 {line_no} 行必须是 引文=编号")
            citation, number = (part.strip() for part in line.split("="))
            if not citation or not number.isdigit() or int(number) < 1:
                raise ValueError(f"第 {line_no} 行的引文或编号无效")
            if citation in mapping:
                raise ValueError(f"第 {line_no} 行的引文重复：{citation}")
            mapping[citation] = int(number)
    if not mapping:
        raise ValueError("映射文件没有有效条目")
    return mapping


def find_reference_section(document):
    """定位参考文献标题及编号条目，要求每个编号唯一。"""
    title_index = next((i for i, p in enumerate(document.paragraphs) if p.text.strip() in REF_TITLES), None)
    if title_index is None:
        raise ValueError("未找到参考文献标题")
    refs = {}
    for index, paragraph in enumerate(document.paragraphs[title_index + 1 :], title_index + 1):
        text = paragraph.text.strip()
        if not text:
            continue
        match = re.match(r"^\[(\d+)\]", text)
        if not match:
            continue
        number = int(match.group(1))
        if number in refs:
            raise ValueError(f"参考文献编号重复：[{number}]")
        refs[number] = index
    if not refs:
        raise ValueError("参考文献区未找到以 [N] 开头的条目")
    return title_index, refs


def make_run(text, source_rpr):
    """沿用原 run 的字符格式，避免引文替换改变正文外观。"""
    run = OxmlElement("w:r")
    if source_rpr is not None:
        run.append(deepcopy(source_rpr))
    node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    return run


def make_ref_runs(bookmark, display, source_rpr):
    """创建 begin、指令、结果、end 四段 Word REF 域，并保留可见文字格式。"""
    result = []
    for kind, value in (
        ("begin", None),
        ("instruction", f" REF {bookmark} \\h \\* MERGEFORMAT"),
        ("separate", None),
        ("display", display),
        ("end", None),
    ):
        run = OxmlElement("w:r")
        if source_rpr is not None:
            run.append(deepcopy(source_rpr))
        if kind in {"begin", "separate", "end"}:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        elif kind == "instruction":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = value
        else:
            node = OxmlElement("w:t")
            node.text = value
        run.append(node)
        result.append(run)
    return result


def add_bookmark(paragraph, number, used_names, used_ids):
    """为文后条目加唯一书签；名称或 ID 冲突时明确报错。"""
    name = f"_AYRef{number}"
    bookmark_id = str(9000 + number)
    if name in used_names or bookmark_id in used_ids:
        raise ValueError(f"书签冲突：{name} 或 id={bookmark_id}")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    return name


def existing_bookmarks(document):
    """收集已有书签，避免覆盖其他 Word 域或用户书签。"""
    names, ids = set(), set()
    for paragraph in document.paragraphs:
        for node in paragraph._p.iter(f"{{{NS}}}bookmarkStart"):
            names.add(node.get(qn("w:name")))
            ids.add(node.get(qn("w:id")))
    return names, ids


def replace_in_run(run, mapping, bookmarks):
    """只替换同一 Word run 内的完整引文，拒绝跨 run 的静默降级。"""
    texts = run.findall(f"{{{NS}}}t")
    if len(texts) != 1:
        return False
    original = texts[0].text or ""
    keys = sorted(mapping, key=len, reverse=True)
    pattern = re.compile("|".join(re.escape(key) for key in keys))
    matches = list(pattern.finditer(original))
    if not matches:
        return False
    parent = run.getparent()
    source_rpr = run.find(f"{{{NS}}}rPr")
    insert_at = parent.index(run)
    cursor = 0
    for match in matches:
        if match.start() > cursor:
            parent.insert(insert_at, make_run(original[cursor : match.start()], source_rpr))
            insert_at += 1
        citation = match.group(0)
        for field_run in make_ref_runs(bookmarks[mapping[citation]], citation, source_rpr):
            parent.insert(insert_at, field_run)
            insert_at += 1
        cursor = match.end()
    if cursor < len(original):
        parent.insert(insert_at, make_run(original[cursor:], source_rpr))
    parent.remove(run)
    return True


def check_cross_run_citations(document, title_index, mapping):
    """当引文仅在段落全文中存在而不在单个 run 中存在时，提示用户先合并该引文。"""
    for index, paragraph in enumerate(document.paragraphs[:title_index]):
        full_text = paragraph.text
        runs = list(paragraph._p.findall(f"{{{NS}}}r"))
        run_texts = ["".join(t.text or "" for t in run.findall(f"{{{NS}}}t")) for run in runs]
        for citation in mapping:
            if citation in full_text and not any(citation in text for text in run_texts):
                raise ValueError(f"第 {index + 1} 段的引文跨 Word run：{citation}；请在 Word 中合并该引文的字符格式后重试")
            for run, text in zip(runs, run_texts):
                if citation in text and len(run.findall(f"{{{NS}}}t")) != 1:
                    raise ValueError(f"第 {index + 1} 段的引文含多个文本节点：{citation}；请在 Word 中合并该引文的字符格式后重试")


def main():
    parser = argparse.ArgumentParser(description="将精确著者—年份引文转换为 Word REF 交叉引用")
    parser.add_argument("input", help="输入 DOCX")
    parser.add_argument("--map-file", required=True, help="UTF-8 映射文件：每行 引文=编号")
    parser.add_argument("--check", action="store_true", help="仅校验映射、编号及引文位置")
    parser.add_argument("--verify", action="store_true", help="写入后校验 REF 域数量")
    args = parser.parse_args()

    mapping = read_mapping(args.map_file)
    document = Document(args.input)
    title_index, refs = find_reference_section(document)
    missing = sorted(set(mapping.values()) - set(refs))
    if missing:
        raise ValueError(f"映射引用了不存在的参考文献编号：{missing}")
    check_cross_run_citations(document, title_index, mapping)
    occurrences = {citation: sum(paragraph.text.count(citation) for paragraph in document.paragraphs[:title_index]) for citation in mapping}
    absent = [citation for citation, count in occurrences.items() if count == 0]
    if absent:
        raise ValueError(f"正文未找到映射引文：{', '.join(absent)}")
    print("映射校验通过：")
    for citation, number in mapping.items():
        print(f"  {citation} -> [{number}]，出现 {occurrences[citation]} 次")
    if args.check:
        return

    backup = os.path.splitext(args.input)[0] + ".bak.docx"
    if not os.path.exists(backup):
        shutil.copy2(args.input, backup)
        print(f"已创建备份：{backup}")
    used_names, used_ids = existing_bookmarks(document)
    bookmarks = {number: add_bookmark(document.paragraphs[index], number, used_names, used_ids) for number, index in refs.items() if number in set(mapping.values())}
    replaced = 0
    for paragraph in document.paragraphs[:title_index]:
        for run in list(paragraph._p.findall(f"{{{NS}}}r")):
            if replace_in_run(run, mapping, bookmarks):
                replaced += 1
    output = os.path.splitext(args.input)[0] + "_著者年份交叉引用.docx"
    document.save(output)
    if args.verify:
        verified = Document(output)
        ref_fields = sum(
            1
            for paragraph in verified.paragraphs
            for node in paragraph._p.iter(f"{{{NS}}}instrText")
            if node.text and " REF _AYRef" in node.text
        )
        if ref_fields != sum(occurrences.values()):
            raise ValueError(f"REF 域数量不一致：预期 {sum(occurrences.values())}，实际 {ref_fields}")
    print(f"已替换 {replaced} 个 run，输出：{output}")


if __name__ == "__main__":
    try:
        main()
    except (OSError, ValueError) as error:
        print(f"错误：{error}", file=sys.stderr)
        sys.exit(1)
