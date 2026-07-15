#!/usr/bin/env python3
"""将 DOCX 顺序编码引文转换为可维护的 Word REF 交叉引用。"""

import argparse
import os
import re
import sys
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REF_TITLES = {"参考文献", "〖参考文献〗", "主要参考文献", "References", "Bibliography"}
CITATION_RE = re.compile(r"\[(\d+(?:\s*[,，、\-–—]\s*\d+)*)\]")


def paragraph_text(paragraph):
    return paragraph.text.strip()


def find_reference_title(document):
    """定位参考文献标题，避免把正文编号误判为文后条目。"""
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph_text(paragraph) in REF_TITLES:
            return index
    raise ValueError("未找到参考文献标题")


def remove_auto_numbering(paragraph):
    """移除 Word 自动编号，只保留段落里的真实文本编号。"""
    ppr = paragraph._p.pPr
    if ppr is None:
        return False
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    ppr.remove(num_pr)
    return True


def first_run_rpr(paragraph):
    first_run = paragraph._p.find(qn("w:r"))
    if first_run is None:
        return None
    return first_run.find(qn("w:rPr"))


def insert_literal_number(paragraph, number):
    """在没有真实 [N] 文本的条目前插入可选中的编号。"""
    run = OxmlElement("w:r")
    source_rpr = first_run_rpr(paragraph)
    if source_rpr is not None:
        run.append(deepcopy(source_rpr))
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = f"[{number}] "
    run.append(text)
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, run)


def inspect_reference_list(document, title_index):
    """统计文后条目的编号形态，供处理前确认当前所处的维护阶段。"""
    total = 0
    literal = 0
    automatic = 0
    for paragraph in document.paragraphs[title_index + 1 :]:
        if not paragraph_text(paragraph):
            continue
        total += 1
        if re.match(r"^\[(\d+)\]", paragraph_text(paragraph)):
            literal += 1
        if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn("w:numPr")) is not None:
            automatic += 1
    return total, literal, automatic


def normalize_reference_list(document, title_index, insert_missing, remove_automatic):
    """统一文后条目：按显式授权去掉自动编号，并按需要补入正文字符编号。"""
    references = {}
    removed_auto = 0
    inserted = 0
    next_number = 1
    for index, paragraph in enumerate(document.paragraphs[title_index + 1 :], title_index + 1):
        text = paragraph_text(paragraph)
        if not text:
            continue
        if remove_automatic and remove_auto_numbering(paragraph):
            removed_auto += 1
        match = re.match(r"^\[(\d+)\]", paragraph_text(paragraph))
        if match is None:
            if not insert_missing:
                continue
            insert_literal_number(paragraph, next_number)
            inserted += 1
            number = next_number
        else:
            number = int(match.group(1))
        if number in references:
            raise ValueError(f"参考文献编号重复：[{number}]")
        references[number] = index
        next_number = number + 1
    if not references:
        raise ValueError("参考文献区没有可识别的条目；如需把自动编号转成文本，请加 --insert-missing-reference-numbers")
    return references, removed_auto, inserted


def existing_bookmarks(document):
    """收集已有书签，避免覆盖用户手工维护的交叉引用目标。"""
    names, ids = set(), set()
    for paragraph in document.paragraphs:
        for bookmark in paragraph._p.iter(qn("w:bookmarkStart")):
            names.add(bookmark.get(qn("w:name")))
            ids.add(bookmark.get(qn("w:id")))
    return names, ids


def bookmark_name(prefix, number):
    return f"{prefix}_{number:03d}"


def add_bookmark(paragraph, number, prefix, names, ids):
    """给参考文献条目添加 Word 界面可见的书签。"""
    name = bookmark_name(prefix, number)
    bookmark_id = str(8000 + number)
    if name.startswith("_"):
        raise ValueError("书签前缀不能以下划线开头；下划线开头会变成 Word 隐藏书签")
    if name in names or bookmark_id in ids:
        raise ValueError(f"书签冲突：{name} 或 id={bookmark_id}")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    names.add(name)
    ids.add(bookmark_id)
    return name


def make_citation_rpr(source_rpr):
    """生成黑色 9 磅上标，符合中文期刊常见顺序编码显示。"""
    rpr = deepcopy(source_rpr) if source_rpr is not None else OxmlElement("w:rPr")
    for tag in ("vertAlign", "sz", "szCs", "color"):
        for node in rpr.findall(qn(f"w:{tag}")):
            rpr.remove(node)
    vertical = OxmlElement("w:vertAlign")
    vertical.set(qn("w:val"), "superscript")
    rpr.append(vertical)
    for tag in ("sz", "szCs"):
        size = OxmlElement(f"w:{tag}")
        size.set(qn("w:val"), "18")
        rpr.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)
    return rpr


def make_text_run(text, source_rpr, superscript=False):
    run = OxmlElement("w:r")
    if superscript:
        run.append(make_citation_rpr(source_rpr))
    elif source_rpr is not None:
        run.append(deepcopy(source_rpr))
    node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    return run


def make_ref_field(bookmark, display, source_rpr):
    """创建显示为上标数字的 REF 域，正文可 Ctrl+单击跳转。"""
    result = []
    for kind, value in (
        ("begin", None),
        ("instruction", f" REF {bookmark} \\h \\* MERGEFORMAT"),
        ("separate", None),
        ("display", display),
        ("end", None),
    ):
        run = OxmlElement("w:r")
        run.append(make_citation_rpr(source_rpr))
        if kind in {"begin", "separate", "end"}:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        elif kind == "instruction":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = value
        else:
            node = OxmlElement("w:t")
            node.text = value
        run.append(node)
        result.append(run)
    return result


def emit_cluster(cluster, bookmarks, source_rpr):
    """保留 [3,27]、[13-18] 的符号，并让可见数字各自可跳转。"""
    pieces = []
    cursor = 0
    for match in re.finditer(r"\d+", cluster):
        if match.start() > cursor:
            pieces.append(make_text_run(cluster[cursor : match.start()], source_rpr, superscript=True))
        number = int(match.group(0))
        if number not in bookmarks:
            raise ValueError(f"正文引用了不存在的参考文献编号：[{number}]")
        pieces.extend(make_ref_field(bookmarks[number], match.group(0), source_rpr))
        cursor = match.end()
    if cursor < len(cluster):
        pieces.append(make_text_run(cluster[cursor:], source_rpr, superscript=True))
    return pieces


def replace_run(run, bookmarks):
    """仅处理单文本节点 run；跨 run 引文必须先在 Word 中合并格式。"""
    texts = run.findall(qn("w:t"))
    if len(texts) != 1:
        return 0
    original = texts[0].text or ""
    matches = list(CITATION_RE.finditer(original))
    if not matches:
        return 0
    parent = run.getparent()
    source_rpr = run.find(qn("w:rPr"))
    insert_at = parent.index(run)
    cursor = 0
    total_numbers = 0
    for match in matches:
        if match.start() > cursor:
            parent.insert(insert_at, make_text_run(original[cursor : match.start()], source_rpr))
            insert_at += 1
        for field_run in emit_cluster(match.group(0), bookmarks, source_rpr):
            parent.insert(insert_at, field_run)
            insert_at += 1
        total_numbers += len(re.findall(r"\d+", match.group(0)))
        cursor = match.end()
    if cursor < len(original):
        parent.insert(insert_at, make_text_run(original[cursor:], source_rpr))
    parent.remove(run)
    return total_numbers


def ensure_no_cross_run_citations(document, title_index):
    """发现跨 run 编号时直接报错，避免生成半断开的引用。"""
    for index, paragraph in enumerate(document.paragraphs[:title_index]):
        full_text = paragraph.text
        run_texts = [
            "".join(text.text or "" for text in run.findall(qn("w:t")))
            for run in paragraph._p.findall(qn("w:r"))
        ]
        for match in CITATION_RE.finditer(full_text):
            citation = match.group(0)
            if not any(citation in text for text in run_texts):
                raise ValueError(f"第 {index + 1} 段的引文跨 Word run：{citation}；请先在 Word 中合并该引文格式")


def main():
    parser = argparse.ArgumentParser(description="将顺序编码引文写为 Word REF 交叉引用")
    parser.add_argument("input", help="输入 DOCX")
    parser.add_argument("--output", help="输出 DOCX；默认在输入文件名后加 _交叉引用")
    parser.add_argument("--normalize-reference-list", action="store_true", help="删除参考文献段落的 Word 自动编号")
    parser.add_argument("--insert-missing-reference-numbers", action="store_true", help="对没有 [N] 文本的文后条目按顺序插入编号")
    parser.add_argument("--bookmark-prefix", default="Ref", help="可见书签前缀，默认 Ref，生成 Ref_001")
    parser.add_argument("--check", action="store_true", help="仅报告文后编号类型，不写入文件")
    parser.add_argument("--verify", action="store_true", help="校验 REF 域数量、书签名称和上标格式")
    args = parser.parse_args()

    document = Document(args.input)
    title_index = find_reference_title(document)
    total, literal, automatic = inspect_reference_list(document, title_index)
    if args.check:
        print(
            f"文后条目 {total} 条；正文字符编号 {literal} 条；"
            f"Word 自动编号 {automatic} 条。"
        )
        if automatic:
            print("当前为写作阶段或混合编号状态；定稿转换请加 --normalize-reference-list。")
        else:
            print("当前文后编号为正文字符编号，可直接用于定稿交叉引用。")
        return
    if automatic and not args.normalize_reference_list:
        raise ValueError("参考文献区存在 Word 自动编号；请加 --normalize-reference-list 后再生成交叉引用")
    references, removed_auto, inserted = normalize_reference_list(
        document,
        title_index,
        args.insert_missing_reference_numbers,
        args.normalize_reference_list,
    )
    ensure_no_cross_run_citations(document, title_index)

    used_numbers = set()
    for paragraph in document.paragraphs[:title_index]:
        for cluster in CITATION_RE.findall(paragraph.text):
            used_numbers.update(int(number) for number in re.findall(r"\d+", cluster))
    if not used_numbers:
        raise ValueError("正文未找到 [N]、[N,N] 或 [N-N] 引文")
    missing = sorted(used_numbers - set(references))
    if missing:
        raise ValueError(f"正文引用了不存在的编号：{missing}")

    names, ids = existing_bookmarks(document)
    bookmarks = {
        number: add_bookmark(document.paragraphs[references[number]], number, args.bookmark_prefix, names, ids)
        for number in sorted(used_numbers)
    }
    fields_written = 0
    for paragraph in document.paragraphs[:title_index]:
        for run in list(paragraph._p.findall(qn("w:r"))):
            fields_written += replace_run(run, bookmarks)

    output = args.output or os.path.splitext(args.input)[0] + "_交叉引用.docx"
    document.save(output)
    if args.verify:
        verified = Document(output)
        field_count = 0
        superscript_count = 0
        for paragraph in verified.paragraphs:
            for node in paragraph._p.iter(qn("w:instrText")):
                if node.text and f" REF {args.bookmark_prefix}_" in node.text:
                    field_count += 1
            for run in paragraph._p.iter(qn("w:r")):
                rpr = run.find(qn("w:rPr"))
                if rpr is not None and rpr.find(qn("w:vertAlign")) is not None:
                    superscript_count += 1
        if field_count != fields_written:
            raise ValueError(f"REF 域数量不一致：预期 {fields_written}，实际 {field_count}")
        if superscript_count < fields_written:
            raise ValueError("上标格式数量异常，请检查输出文档")
    print(f"已写入 {fields_written} 个 REF 域；删除自动编号 {removed_auto} 处；补入文本编号 {inserted} 处；输出：{output}")


if __name__ == "__main__":
    try:
        main()
    except (OSError, ValueError) as error:
        print(f"错误：{error}", file=sys.stderr)
        sys.exit(1)
