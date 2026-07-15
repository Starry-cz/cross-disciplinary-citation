#!/usr/bin/env python3
"""将文后真实的 [N] 编号转为 Word 自动编号，供论文写作阶段使用。"""

import argparse
import os
import re
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REF_TITLES = {"参考文献", "〖参考文献〗", "主要参考文献", "References", "Bibliography"}
LITERAL_NUMBER_RE = re.compile(r"^\[(\d+)\][ \t]*")


def find_reference_title(document):
    """定位文后参考文献标题，避免误处理正文中的方括号引文。"""
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() in REF_TITLES:
            return index
    raise ValueError("未找到参考文献标题")


def next_identifier(nodes, attribute):
    """为 numbering.xml 生成未占用的数值标识。"""
    values = [int(node.get(qn(attribute))) for node in nodes if node.get(qn(attribute), "").isdigit()]
    return max(values, default=0) + 1


def add_bracket_numbering(document):
    """新增形如 [1] 的单级 Word 列表定义，并返回 numId。"""
    numbering = document.part.numbering_part._element
    abstract_id = next_identifier(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    num_id = next_identifier(numbering.findall(qn("w:num")), "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    number_text = OxmlElement("w:lvlText")
    number_text.set(qn("w:val"), "[%1]")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "360")
    indent.set(qn("w:hanging"), "360")
    paragraph_properties.extend([tabs, indent])
    level.extend([start, number_format, number_text, suffix, justification, paragraph_properties])
    abstract.append(level)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.extend([abstract, number])
    return num_id


def remove_literal_number(paragraph, expected_number):
    """删除段首 [N] 文本；编号不连续时停止，避免错误重排用户文献。"""
    match = LITERAL_NUMBER_RE.match(paragraph.text)
    if match is None:
        raise ValueError(f"参考文献缺少正文字符编号：{paragraph.text[:60]}")
    actual_number = int(match.group(1))
    if actual_number != expected_number:
        raise ValueError(f"参考文献编号不连续：预期 [{expected_number}]，实际 [{actual_number}]")

    remaining = len(match.group(0))
    for run in paragraph.runs:
        if remaining == 0:
            break
        text = run.text
        if not text:
            continue
        removed = min(len(text), remaining)
        run.text = text[removed:]
        remaining -= removed
    if remaining:
        raise ValueError("段首编号跨越了不支持的 Word 域或对象，未执行转换")


def set_auto_numbering(paragraph, num_id):
    """将段落绑定至本次新增的 Word 自动编号列表。"""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    old_numbering = paragraph_properties.find(qn("w:numPr"))
    if old_numbering is not None:
        paragraph_properties.remove(old_numbering)
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.extend([level, number])
    paragraph_properties.append(number_properties)


def main():
    parser = argparse.ArgumentParser(description="将参考文献 [N] 转为 Word 自动编号")
    parser.add_argument("input", help="输入 DOCX")
    parser.add_argument("--output", help="输出 DOCX；默认在输入文件名后加 _自动编号")
    parser.add_argument("--verify", action="store_true", help="校验自动编号段落数量与文字编号移除结果")
    args = parser.parse_args()

    document = Document(args.input)
    title_index = find_reference_title(document)
    reference_paragraphs = [
        paragraph
        for paragraph in document.paragraphs[title_index + 1 :]
        if LITERAL_NUMBER_RE.match(paragraph.text)
    ]
    if not reference_paragraphs:
        raise ValueError("参考文献区未找到可转换的 [N] 正文字符编号")

    num_id = add_bracket_numbering(document)
    for expected_number, paragraph in enumerate(reference_paragraphs, start=1):
        remove_literal_number(paragraph, expected_number)
        set_auto_numbering(paragraph, num_id)

    output = args.output or os.path.splitext(args.input)[0] + "_自动编号.docx"
    document.save(output)
    if args.verify:
        verified = Document(output)
        verified_title = find_reference_title(verified)
        verified_paragraphs = [
            paragraph
            for paragraph in verified.paragraphs[verified_title + 1 :]
            if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn("w:numPr")) is not None
        ]
        remaining_literals = [
            paragraph.text
            for paragraph in verified.paragraphs[verified_title + 1 :]
            if LITERAL_NUMBER_RE.match(paragraph.text)
        ]
        if len(verified_paragraphs) != len(reference_paragraphs):
            raise ValueError("Word 自动编号段落数量异常")
        if remaining_literals:
            raise ValueError("仍存在未移除的 [N] 正文字符编号")
    print(f"已将 {len(reference_paragraphs)} 条参考文献转换为 Word 自动编号；输出：{output}")


if __name__ == "__main__":
    try:
        main()
    except (OSError, ValueError) as error:
        print(f"错误：{error}", file=sys.stderr)
        sys.exit(1)
